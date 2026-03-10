"""
Resume Parser — Two-Layer Architecture
Layer 1: Raw text extraction (PyMuPDF → pdfminer fallback for PDF; python-docx for DOCX)
Layer 2: Structured field extraction (spaCy + regex)
"""
import re
import io
from pathlib import Path
from typing import Tuple, Dict, Any


# ──────────────────────────────────────────────────
#  LAYER 1 — RAW TEXT EXTRACTION
# ──────────────────────────────────────────────────

def extract_text_pdf(file_bytes: bytes) -> str:
    """Extract raw text from PDF using PyMuPDF; falls back to pdfminer."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = [page.get_text("text") for page in doc]
        text = "\n".join(pages)
        if text.strip():
            return text
    except Exception:
        pass

    # pdfminer fallback
    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
        output = io.StringIO()
        extract_text_to_fp(io.BytesIO(file_bytes), output, laparams=LAParams())
        text = output.getvalue()
        if text.strip():
            return text
    except Exception:
        pass

    return ""


def extract_text_docx(file_bytes: bytes) -> str:
    """Extract raw text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        return ""


def extract_raw_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the correct parser based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_docx(file_bytes)
    return ""


# ──────────────────────────────────────────────────
#  LAYER 2 — STRUCTURED INFORMATION EXTRACTION
# ──────────────────────────────────────────────────

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(
    r"(\+?\d{1,3}[\s\-]?)?(\(?\d{2,4}\)?[\s\-]?)?\d{3,4}[\s\-]?\d{3,4}"
)

SECTION_HEADERS = {
    "education": r"(education|academic|qualification|degree|university|college)",
    "experience": r"(experience|work history|employment|professional background|career)",
    "skills": r"(skills|technologies|tech stack|competencies|expertise|proficiencies)",
    "certifications": r"(certification|certificate|credential|license|award)",
}


def _extract_email(text: str):
    match = EMAIL_RE.search(text)
    return match.group(0) if match else None


def _extract_phone(text: str):
    match = PHONE_RE.search(text)
    return match.group(0).strip() if match else None


def _extract_name_spacy(text: str):
    """Use spaCy PERSON entity from first 300 chars for name extraction."""
    try:
        from app.ml.preprocessing import get_nlp
        nlp = get_nlp()
        snippet = text[:300]
        doc = nlp(snippet)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                return ent.text.strip()
    except Exception:
        pass
    # Fallback: first non-empty line
    for line in text.splitlines():
        line = line.strip()
        if line and len(line.split()) <= 5 and not EMAIL_RE.search(line):
            return line
    return None


def _extract_section(text: str, section_key: str) -> list:
    pattern = SECTION_HEADERS[section_key]
    lines = text.splitlines()
    result = []
    in_section = False

    other_headers = "|".join(
        v for k, v in SECTION_HEADERS.items() if k != section_key
    )

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if re.search(pattern, stripped, re.IGNORECASE):
            in_section = True
            continue

        if in_section:
            if re.search(other_headers, stripped, re.IGNORECASE):
                break
            if len(stripped) > 2:
                result.append(stripped)

        if len(result) >= 20:
            break

    return result

def parse_resume(file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
    """
    Returns:
      - raw_text: original unmodified text
      - parsed: dict with structured fields
    """
    raw_text = extract_raw_text(file_bytes, filename)

    parsed = {
        "name": _extract_name_spacy(raw_text),
        "email": _extract_email(raw_text),
        "phone": _extract_phone(raw_text),
        "skills": [],          # Filled by skill_extractor
        "education": _extract_section(raw_text, "education"),
        "experience": _extract_section(raw_text, "experience"),
        "certifications": _extract_section(raw_text, "certifications"),
    }

    return raw_text, parsed
